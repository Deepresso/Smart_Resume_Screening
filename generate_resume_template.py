"""
Run this script once to generate the resume template Word file.
Output: static/downloads/resume_template.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    set_font(run, size=10, bold=True, color=(234, 88, 12))
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'EA580C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, indent=0.3):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p


doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin   = Inches(0.85)
    section.right_margin  = Inches(0.85)

# ── TIP BANNER ──
tip = doc.add_paragraph()
tip.paragraph_format.space_after = Pt(10)
tip_run = tip.add_run(
    '📋  HOW TO USE THIS TEMPLATE\n'
    'Replace every [bracketed placeholder] with your own information. '
    'Keep the section headings and structure as-is — the screening system scores '
    'your resume against the job\'s keywords and description, so include relevant '
    'skills and experience using clear, specific language. '
    'Save your completed resume as PDF or DOCX before uploading.'
)
set_font(tip_run, size=9.5, italic=True, color=(107, 114, 128))

# ── NAME ──
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(2)
name_run = name_p.add_run('[YOUR FULL NAME]')
set_font(name_run, size=20, bold=True, color=(30, 41, 59))

# ── CONTACT LINE ──
contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(4)
contact_run = contact_p.add_run(
    '[Phone Number]  ·  [Email Address]  ·  [LinkedIn URL]  ·  [City, State]'
)
set_font(contact_run, size=10, color=(107, 114, 128))

# ── PROFESSIONAL SUMMARY ──
add_section_heading(doc, 'Professional Summary')
summary = doc.add_paragraph()
summary.paragraph_format.space_after = Pt(4)
summary_run = summary.add_run(
    'Results-driven [Job Title] with [X] years of experience in [your industry/domain]. '
    'Proficient in [Skill 1], [Skill 2], and [Skill 3]. Demonstrated ability to [key achievement]. '
    'Seeking to contribute technical expertise and problem-solving skills to a dynamic team.'
)
set_font(summary_run, size=10.5)

# ── TECHNICAL SKILLS ──
add_section_heading(doc, 'Technical Skills')

skills_table = doc.add_table(rows=4, cols=2)
skills_table.style = 'Table Grid'
skills_data = [
    ('Programming Languages', '[Python / Java / JavaScript / C++ / etc.]'),
    ('Frameworks & Libraries', '[React / Flask / Django / Node.js / etc.]'),
    ('Databases',              '[MySQL / PostgreSQL / MongoDB / SQLite / etc.]'),
    ('Tools & Platforms',      '[Git / Docker / AWS / Linux / Jira / etc.]'),
]
for i, (label, value) in enumerate(skills_data):
    row = skills_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, size=10.5)
        cell.width = Inches(2.5)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── WORK EXPERIENCE ──
add_section_heading(doc, 'Work Experience')

# Job 1
exp1_title = doc.add_paragraph()
exp1_title.paragraph_format.space_after = Pt(0)
r1 = exp1_title.add_run('[Job Title]')
set_font(r1, size=11, bold=True)

exp1_meta = doc.add_paragraph()
exp1_meta.paragraph_format.space_after = Pt(3)
r1m = exp1_meta.add_run('[Company Name]  ·  [City, State]  ·  [Month Year] – [Month Year / Present]')
set_font(r1m, size=10, italic=True, color=(107, 114, 128))

add_bullet(doc, 'Developed and maintained [type of application/system] using [technologies], resulting in [measurable outcome].')
add_bullet(doc, 'Collaborated with [team/stakeholders] to design and implement [feature/module], improving [metric] by [X]%.')
add_bullet(doc, 'Responsible for [task], utilising [tools/frameworks] to [outcome].')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Job 2
exp2_title = doc.add_paragraph()
exp2_title.paragraph_format.space_after = Pt(0)
r2 = exp2_title.add_run('[Previous Job Title]')
set_font(r2, size=11, bold=True)

exp2_meta = doc.add_paragraph()
exp2_meta.paragraph_format.space_after = Pt(3)
r2m = exp2_meta.add_run('[Company Name]  ·  [City, State]  ·  [Month Year] – [Month Year]')
set_font(r2m, size=10, italic=True, color=(107, 114, 128))

add_bullet(doc, 'Designed and deployed [system/feature] that [outcome/impact].')
add_bullet(doc, 'Wrote unit and integration tests using [testing framework], achieving [X]% code coverage.')
add_bullet(doc, 'Participated in code reviews and contributed to team knowledge sharing sessions.')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── EDUCATION ──
add_section_heading(doc, 'Education')

edu_title = doc.add_paragraph()
edu_title.paragraph_format.space_after = Pt(0)
re1 = edu_title.add_run('[Bachelor\'s / Diploma] in [Field of Study]')
set_font(re1, size=11, bold=True)

edu_meta = doc.add_paragraph()
edu_meta.paragraph_format.space_after = Pt(3)
rem = edu_meta.add_run('[University Name]  ·  [City, State]  ·  Graduated [Month Year]  ·  CGPA: [X.XX / 4.00]')
set_font(rem, size=10, italic=True, color=(107, 114, 128))

add_bullet(doc, 'Relevant coursework: [Course 1], [Course 2], [Course 3], [Course 4].')
add_bullet(doc, '[Academic achievement, scholarship, or dean\'s list mention if applicable].')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── PROJECTS ──
add_section_heading(doc, 'Projects')

proj1_title = doc.add_paragraph()
proj1_title.paragraph_format.space_after = Pt(0)
rp1 = proj1_title.add_run('[Project Name]')
set_font(rp1, size=11, bold=True)

proj1_tech = doc.add_paragraph()
proj1_tech.paragraph_format.space_after = Pt(3)
rp1t = proj1_tech.add_run('Technologies: [Tech Stack Used]')
set_font(rp1t, size=10, italic=True, color=(107, 114, 128))

add_bullet(doc, 'Built [what the project does] using [technologies], which [outcome/purpose].')
add_bullet(doc, 'Implemented [key feature] that [benefit/impact].')
add_bullet(doc, 'GitHub / Live URL: [link if applicable]')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── CERTIFICATIONS ──
add_section_heading(doc, 'Certifications & Training')

add_bullet(doc, '[Certification Name] — [Issuing Organisation] — [Year]')
add_bullet(doc, '[Online Course / Training Name] — [Platform, e.g. Coursera / Udemy] — [Year]')

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── SCORING TIPS ──
add_section_heading(doc, 'Resume Scoring Tips  (Read before submitting — then delete this section)')

tips = [
    'Match the job keywords: Check the "Screening Keywords" listed on the job posting page and make sure those exact words appear in your resume.',
    'Use specific technology names: Write "Python", "MySQL", "REST API" — not just "programming" or "databases".',
    'Describe impact with numbers: "Improved load time by 40%" scores better than "improved performance".',
    'Fill in every section: A complete resume (summary + skills + experience + education + projects) scores significantly higher than a short one.',
    'Avoid tables and graphics for skills: The system reads plain text — skills listed in a table may not be extracted properly. List them as plain text too.',
    'Save as PDF: PDF preserves formatting and extracts cleanly. DOCX also works.',
]
for tip in tips:
    add_bullet(doc, tip)

# ── Save ──
out = 'static/downloads/resume_template.docx'
doc.save(out)
print(f'Template saved to {out}')
