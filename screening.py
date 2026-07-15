import re
import numpy as np
import fitz  # PyMuPDF
from docx import Document
from rapidfuzz import fuzz
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Lazy-loaded BERT model — downloaded on first use (~80 MB)
_bert_model = None

def _get_bert_model():
    global _bert_model
    if _bert_model is None:
        from sentence_transformers import SentenceTransformer
        _bert_model = SentenceTransformer('all-MiniLM-L6-v2')
    return _bert_model


# ── Text Extraction ──────────────────────────────────────────────────────────

def extract_text(filepath):
    if filepath.lower().endswith('.pdf'):
        return _from_pdf(filepath)
    elif filepath.lower().endswith('.docx'):
        return _from_docx(filepath)
    return ''


def _from_pdf(filepath):
    text = ''
    with fitz.open(filepath) as doc:
        for page in doc:
            text += page.get_text()
    return text.strip()


def _from_docx(filepath):
    doc = Document(filepath)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
    return '\n'.join(parts).strip()


# ── NLP Scoring ──────────────────────────────────────────────────────────────

def keyword_score(resume_text, keywords):
    """Percentage of job keywords found in resume (case-insensitive)."""
    if not keywords:
        return 0.0
    resume_lower = resume_text.lower()
    matched = sum(1 for kw in keywords if kw.lower() in resume_lower)
    return round((matched / len(keywords)) * 100, 2)


def fuzzy_score(resume_text, keywords):
    """Average best fuzzy match score for each keyword against resume words."""
    if not keywords:
        return 0.0
    resume_words = re.findall(r'\b\w+\b', resume_text.lower())
    scores = []
    for kw in keywords:
        best = max((fuzz.ratio(kw.lower(), word) for word in resume_words), default=0)
        scores.append(best)
    return round(sum(scores) / len(scores), 2)


def similarity_score(resume_text, job_description):
    """TF-IDF cosine similarity between resume and job description."""
    if not resume_text or not job_description:
        return 0.0
    vectorizer = TfidfVectorizer(stop_words='english')
    try:
        tfidf = vectorizer.fit_transform([resume_text, job_description])
        score = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
        return round(float(score) * 100, 2)
    except Exception:
        return 0.0


def semantic_score(resume_text, job_description):
    """BERT cosine similarity between resume and job description embeddings."""
    if not resume_text or not job_description:
        return 0.0
    try:
        model = _get_bert_model()
        # Truncate to ~1000 chars — BERT has a 512-token limit
        emb = model.encode([resume_text[:1000], job_description[:1000]])
        a, b = emb[0], emb[1]
        score = float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b)))
        return round(max(score, 0.0) * 100, 2)
    except Exception:
        return 0.0


def keyword_breakdown(resume_text, keywords):
    """Per-keyword found/not-found list."""
    resume_lower = resume_text.lower()
    return [{'kw': kw, 'found': kw.lower() in resume_lower} for kw in keywords]


def fuzzy_breakdown(resume_text, keywords):
    """Per-keyword best fuzzy match score (0-100)."""
    resume_words = re.findall(r'\b\w+\b', resume_text.lower())
    result = []
    for kw in keywords:
        best = max((fuzz.ratio(kw.lower(), word) for word in resume_words), default=0)
        result.append({'kw': kw, 'score': round(best, 1)})
    return result


def compute_scores(resume_text, job_description, keywords):
    """Run four NLP algorithms and return individual + composite scores."""
    kw_list = [k.strip() for k in keywords.split(',') if k.strip()] if keywords else []

    kw  = keyword_score(resume_text, kw_list)
    fz  = fuzzy_score(resume_text, kw_list)
    sim = similarity_score(resume_text, job_description)
    sem = semantic_score(resume_text, job_description)

    # Weighted composite: keyword 40%, fuzzy 20%, TF-IDF 10%, BERT 30%
    composite = round((kw * 0.40) + (fz * 0.20) + (sim * 0.10) + (sem * 0.30), 2)

    return {
        'keyword_score':    kw,
        'fuzzy_score':      fz,
        'similarity_score': sim,
        'semantic_score':   sem,
        'composite_score':  composite,
    }
