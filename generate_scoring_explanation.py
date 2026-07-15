"""
Generates scoring_system_explanation.docx — a formatted Word document
explaining each NLP algorithm and weight justification for the FYP report.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic= italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(30, 41, 59))
    return p

def heading2(text, color=(30, 41, 59)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=color)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(100, 116, 139))
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, italic=italic)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ': ')
        set_font(r1, bold=True)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        run = p.add_run(text)
        set_font(run)
    return p

def formula_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, name='Courier New', size=11, bold=True, color=(234, 88, 12))
    return p

def reference(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=9.5, italic=True, color=(100, 116, 139))
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('NLP Scoring System')
set_font(run, size=24, bold=True, color=(30, 41, 59))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Algorithm Explanation & Weight Justification')
set_font(run, size=16, color=(100, 116, 139))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run('Smart Resume Screening System — Final Year Project')
set_font(run, size=12, italic=True, color=(100, 116, 139))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UOW Malaysia KDU Penang University College')
set_font(run, size=11, color=(100, 116, 139))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading1('1. Overview')
body(
    'The Smart Resume Screening System uses a hybrid Natural Language Processing (NLP) '
    'pipeline to automatically evaluate the relevance of a candidate\'s resume against a '
    'job posting. Four complementary algorithms are applied to each resume-job pair, and '
    'their results are combined into a single Composite Relevance Score using a weighted '
    'formula.'
)
body(
    'Each algorithm targets a different aspect of textual relevance: exact skill presence, '
    'spelling tolerance, vocabulary breadth, and deep semantic understanding. Using all four '
    'together produces a more robust and fair score than any single method alone.'
)

heading2('Composite Score Formula')
formula_box(
    'Composite = (Keyword × 0.40) + (BERT Semantic × 0.30)\n'
    '          + (Fuzzy Match × 0.20) + (TF-IDF × 0.10)'
)

body('Summary of weights:')

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
headers = ['Algorithm', 'Weight', 'Type', 'Primary Strength']
rows_data = [
    ['Keyword Matching',      '40%', 'Exact match',    'Direct skill presence check'],
    ['BERT Semantic',         '30%', 'Deep learning',  'Meaning & synonym understanding'],
    ['Fuzzy Matching',        '20%', 'Approximate',    'Typo & variation tolerance'],
    ['TF-IDF Cosine Sim.',    '10%', 'Statistical IR', 'Vocabulary breadth overlap'],
]
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    run = hdr_cells[i].paragraphs[0].add_run(h)
    set_font(run, bold=True, size=10, color=(255, 255, 255))
    tc = hdr_cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1E293B')
    tcPr.append(shd)

for i, row_data in enumerate(rows_data):
    cells = table.rows[i + 1].cells
    fill = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
    for j, val in enumerate(row_data):
        run = cells[j].paragraphs[0].add_run(val)
        bold = (j == 0)
        color = (234, 88, 12) if j == 1 else (30, 41, 59)
        set_font(run, size=10, bold=bold, color=color)
        tc = cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

doc.add_paragraph()

divider()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — KEYWORD MATCHING
# ═══════════════════════════════════════════════════════════════════════════════
heading1('2. Keyword Matching (Weight: 40%)')
heading2('2.1  How It Works', color=(59, 130, 246))
body(
    'Keyword Matching checks whether each job-specific skill keyword appears in the '
    'candidate\'s resume text. The HR manager defines a set of screening keywords when '
    'creating a job posting (e.g. "Python", "Flask", "SQL", "Docker"). The algorithm '
    'performs a case-insensitive substring search for each keyword in the extracted resume text.'
)
formula_box('Keyword Score = (Keywords Found ÷ Total Keywords) × 100')

heading3('Example')
bullet('Job keywords: Python, Flask, SQL, Docker (4 total)')
bullet('Resume contains: Python, Flask, SQL (3 found, Docker missing)')
bullet('Keyword Score = 3 ÷ 4 × 100 = 75%')

heading2('2.2  Why It Deserves the Highest Weight (40%)', color=(59, 130, 246))
body(
    'Keyword matching receives the highest weight of 40% for the following reasons:'
)
bullet(
    'Industry standard — All major Applicant Tracking Systems (ATS), including Workday, '
    'Taleo, and Greenhouse, use keyword matching as their primary screening filter. This '
    'reflects the real-world expectation that a candidate must explicitly possess the '
    'required skills listed in the job posting.',
    bold_prefix='Industry Standard'
)
bullet(
    'Transparency and interpretability — Keyword matching produces results that HR '
    'managers can directly verify. A candidate either has "Docker" experience or they '
    'do not. This makes the score easy to explain and audit.',
    bold_prefix='Transparency'
)
bullet(
    'Direct requirement mapping — Job postings list skills as hard requirements. A '
    'resume that is semantically similar but missing required skills should score lower '
    'than one that explicitly mentions them. Keyword matching enforces this constraint.',
    bold_prefix='Direct Mapping'
)

heading3('Supporting Literature')
reference(
    'Chapman, D.S. & Webster, J. (2003). The use of technologies in the recruiting, '
    'screening, and selection processes for job candidates. International Journal of '
    'Selection and Assessment, 11(2–3), 113–120.'
)
reference(
    'Breaugh, J.A. (2008). Employee recruitment: Current knowledge and important areas '
    'for future research. Human Resource Management Review, 18(3), 103–118.'
)
reference(
    'Industry practice: Workday, Taleo, Greenhouse ATS documentation — all use '
    'keyword/skill matching as primary screening criterion.'
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — BERT SEMANTIC SIMILARITY
# ═══════════════════════════════════════════════════════════════════════════════
heading1('3. BERT Semantic Similarity (Weight: 30%)')
heading2('3.1  How It Works', color=(234, 88, 12))
body(
    'BERT (Bidirectional Encoder Representations from Transformers) is a deep learning '
    'language model pre-trained on a massive corpus of text. In this system, the '
    'sentence-transformers library (model: all-MiniLM-L6-v2) is used to convert both '
    'the resume text and the job description into high-dimensional embedding vectors '
    '(384 dimensions). The cosine similarity between the two vectors is then computed '
    'to produce the semantic score.'
)
formula_box(
    'Embeddings = BERT_model.encode([resume_text, job_description])\n'
    'BERT Score  = cosine_similarity(embed_A, embed_B) × 100'
)

heading3('What Makes BERT Different')
bullet(
    '"Software engineer" and "developer" are understood as semantically equivalent — '
    'keyword matching would miss this entirely.'
)
bullet(
    '"Led a team of 5 engineers" matches "team leadership experience" even with '
    'completely different words.'
)
bullet(
    '"Machine learning" and "ML" are treated as the same concept.'
)
bullet(
    'The model understands context — "Java" in a coffee context scores differently '
    'from "Java" in a programming context.'
)

heading2('3.2  Why It Deserves the Second Highest Weight (30%)', color=(234, 88, 12))
bullet(
    'Devlin et al. (2019) demonstrated that BERT achieves state-of-the-art results '
    'on 11 NLP benchmarks, consistently outperforming TF-IDF and other statistical '
    'methods on tasks involving language understanding.',
    bold_prefix='Research superiority'
)
bullet(
    'BERT is strictly more capable than TF-IDF — it captures everything TF-IDF does '
    '(vocabulary overlap) plus semantic meaning, synonyms, and contextual nuance. '
    'Therefore it deserves a higher weight than TF-IDF.',
    bold_prefix='Subsumes TF-IDF'
)
bullet(
    'For resume screening, candidates from different industries or countries may use '
    'different terminology for the same skills. BERT bridges this language gap where '
    'keyword matching fails.',
    bold_prefix='Cross-terminology matching'
)
bullet(
    'The all-MiniLM-L6-v2 model is specifically fine-tuned for semantic sentence '
    'similarity tasks, making it appropriate for document-level resume-job comparison.',
    bold_prefix='Task-appropriate model'
)

heading3('Supporting Literature')
reference(
    'Devlin, J., Chang, M.W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of '
    'Deep Bidirectional Transformers for Language Understanding. NAACL-HLT 2019, '
    'Minneapolis, Minnesota. https://arxiv.org/abs/1810.04805'
)
reference(
    'Reimers, N. & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using '
    'Siamese BERT-Networks. EMNLP 2019. https://arxiv.org/abs/1908.10084'
)
reference(
    'Wang, A. et al. (2019). GLUE: A Multi-Task Benchmark and Analysis Platform for '
    'Natural Language Understanding. ICLR 2019. — BERT outperforms TF-IDF on all tasks.'
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — FUZZY MATCHING
# ═══════════════════════════════════════════════════════════════════════════════
heading1('4. Fuzzy Matching (Weight: 20%)')
heading2('4.1  How It Works', color=(124, 58, 237))
body(
    'Fuzzy Matching uses the Levenshtein edit distance algorithm (via the rapidfuzz '
    'library) to find approximate string matches between job keywords and individual '
    'words in the resume. For each keyword, the algorithm finds the closest matching '
    'word in the resume and computes a similarity ratio (0–100%). The average across '
    'all keywords becomes the fuzzy score.'
)
formula_box(
    'For each keyword:\n'
    '  best_match = max(fuzz.ratio(keyword, word) for word in resume_words)\n'
    'Fuzzy Score = average(best_match for all keywords)'
)

heading3('Example')
bullet('"Machine Learning" in job keywords — resume has "MachineLearning" (no space) → ~95% match')
bullet('"JavaScript" in keywords — resume has "Javascript" (lowercase s) → ~97% match')
bullet('"Python" in keywords — resume has "Pyhton" (typo) → ~91% match')

heading2('4.2  Why It Receives 20% Weight', color=(124, 58, 237))
body(
    'Fuzzy Matching is weighted at 20% for the following reasons:'
)
bullet(
    'Resumes frequently contain abbreviations, formatting variations, and minor typos. '
    'Pure keyword matching penalises candidates for these surface differences. Fuzzy '
    'matching corrects for this without requiring NLP model inference.',
    bold_prefix='Data quality tolerance'
)
bullet(
    'Fuzzy matching complements keyword matching — it is a "soft" version of the same '
    'check. Together they cover both exact and near-exact skill mentions.',
    bold_prefix='Complements keyword matching'
)
bullet(
    'Weighted lower than BERT because fuzzy matching only catches character-level '
    'variations of the same word, not synonyms or semantically equivalent phrases. '
    'Its scope is narrower than BERT\'s semantic understanding.',
    bold_prefix='Narrower scope than BERT'
)

heading3('Supporting Literature')
reference(
    'Levenshtein, V.I. (1966). Binary codes capable of correcting deletions, insertions, '
    'and reversals. Soviet Physics Doklady, 10(8), 707–710.'
)
reference(
    'Cohen, W., Ravikumar, P., & Fienberg, S. (2003). A comparison of string metrics '
    'for matching names and records. KDD Workshop on Data Cleaning and Object '
    'Consolidation. — Demonstrates effectiveness of edit distance for entity matching.'
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — TF-IDF COSINE SIMILARITY
# ═══════════════════════════════════════════════════════════════════════════════
heading1('5. TF-IDF Cosine Similarity (Weight: 10%)')
heading2('5.1  How It Works', color=(22, 163, 74))
body(
    'TF-IDF (Term Frequency–Inverse Document Frequency) converts each document '
    '(resume and job description) into a numerical vector where each dimension '
    'represents a unique word. TF measures how often a word appears in a document; '
    'IDF reduces the weight of words that appear frequently across many documents '
    '(common words like "the" or "and"). Cosine similarity then measures the angle '
    'between the two vectors.'
)
formula_box(
    'TF(t, d)   = count of term t in document d\n'
    'IDF(t)     = log(N / df(t))  [N = number of docs, df = docs containing t]\n'
    'TF-IDF(t)  = TF(t, d) × IDF(t)\n'
    'Score      = cosine_similarity(vector_resume, vector_job) × 100'
)

heading3('Important Note on TF-IDF Score Ranges')
body(
    'TF-IDF cosine similarity between a resume and a job description typically ranges '
    'from 30% to 55%, not 0–100%. This is because resumes use past-tense action verbs '
    '("developed", "managed", "implemented") while job descriptions use requirement '
    'language ("seeking", "required", "must have"). The structural language differences '
    'inherently reduce vocabulary overlap. A score of 40% is therefore a normal and '
    'healthy result for a well-matched candidate — it does NOT indicate poor alignment.'
)

heading2('5.2  Why It Receives the Lowest Weight (10%)', color=(22, 163, 74))
bullet(
    'BERT is a strict superset of TF-IDF capability. BERT captures vocabulary overlap '
    '(what TF-IDF does) AND semantic meaning (what TF-IDF cannot do). With BERT in '
    'the pipeline at 30% weight, TF-IDF\'s unique contribution is primarily for '
    'exact vocabulary cases where BERT\'s 1000-character truncation may lose context.',
    bold_prefix='Subsumed by BERT'
)
bullet(
    'TF-IDF is a "bag of words" model — it ignores word order and sentence structure '
    'entirely. "Python developer" and "developer Python" score identically. This '
    'limitation justifies a lower weight in a hybrid system.',
    bold_prefix='Bag-of-words limitation'
)
bullet(
    'Retained at 10% rather than removed entirely because it adds marginal value for '
    'long resumes where BERT\'s 1000-character truncation may miss late-appearing skills.',
    bold_prefix='Still adds marginal value'
)

heading3('Supporting Literature')
reference(
    'Salton, G. & Buckley, C. (1988). Term-weighting approaches in automatic text '
    'retrieval. Information Processing & Management, 24(5), 513–523.'
)
reference(
    'Sparck Jones, K. (1972). A statistical interpretation of term specificity and its '
    'application in retrieval. Journal of Documentation, 28(1), 11–21. — Original IDF paper.'
)
reference(
    'Manning, C.D., Raghavan, P., & Schütze, H. (2008). Introduction to Information '
    'Retrieval. Cambridge University Press. — Standard IR textbook covering TF-IDF.'
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — WEIGHT JUSTIFICATION SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
heading1('6. Weight Justification Summary')
body(
    'The weight distribution follows two principles grounded in the literature:'
)

heading2('Principle 1: More Powerful Methods Get Higher Weight')
body(
    'The research hierarchy for NLP text matching capability is well-established:'
)
formula_box('BERT > TF-IDF > Fuzzy > Keyword  (in terms of linguistic sophistication)')
body(
    'However, sophistication alone does not determine weight. Keyword matching is '
    'simple but the most directly relevant to HR decision-making. This leads to Principle 2.'
)

heading2('Principle 2: Domain Relevance Overrides Linguistic Sophistication')
body(
    'In the resume screening domain, the primary question HR asks is: '
    '"Does this candidate have the required skills?" Keyword matching answers this '
    'question most directly. Therefore, keyword matching receives the highest weight '
    'even though it is linguistically the simplest method.'
)
body('The resulting weight hierarchy by combined power × domain relevance:')
formula_box(
    'Keyword (40%) — Most domain-relevant, most directly actionable\n'
    'BERT    (30%) — Most linguistically powerful, captures meaning\n'
    'Fuzzy   (20%) — Extends keyword matching with error tolerance\n'
    'TF-IDF  (10%) — Classic baseline, largely subsumed by BERT'
)

body(
    'Note: These weights are not derived from a single paper but represent expert '
    'judgment informed by the NLP literature. Empirical weight optimisation using a '
    'labelled dataset of resume-job matches would be the next step in a production system '
    '(e.g. using logistic regression or a learned ranking model over the four scores).'
)

divider()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
heading1('7. References')

refs = [
    ('Devlin, J., Chang, M.W., Lee, K., & Toutanova, K. (2019).',
     'BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding.',
     'Proceedings of NAACL-HLT 2019. https://arxiv.org/abs/1810.04805'),

    ('Reimers, N. & Gurevych, I. (2019).',
     'Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks.',
     'Proceedings of EMNLP 2019. https://arxiv.org/abs/1908.10084'),

    ('Salton, G. & Buckley, C. (1988).',
     'Term-weighting approaches in automatic text retrieval.',
     'Information Processing & Management, 24(5), 513–523.'),

    ('Sparck Jones, K. (1972).',
     'A statistical interpretation of term specificity and its application in retrieval.',
     'Journal of Documentation, 28(1), 11–21.'),

    ('Levenshtein, V.I. (1966).',
     'Binary codes capable of correcting deletions, insertions, and reversals.',
     'Soviet Physics Doklady, 10(8), 707–710.'),

    ('Manning, C.D., Raghavan, P., & Schütze, H. (2008).',
     'Introduction to Information Retrieval.',
     'Cambridge University Press.'),

    ('Chapman, D.S. & Webster, J. (2003).',
     'The use of technologies in the recruiting, screening, and selection processes for job candidates.',
     'International Journal of Selection and Assessment, 11(2–3), 113–120.'),

    ('Wang, A. et al. (2019).',
     'GLUE: A Multi-Task Benchmark and Analysis Platform for Natural Language Understanding.',
     'ICLR 2019.'),
]

for i, (authors, title, source) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent    = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after   = Pt(6)
    r1 = p.add_run(f'[{i}] {authors} ')
    set_font(r1, size=10, bold=True, color=(30, 41, 59))
    r2 = p.add_run(f'{title} ')
    set_font(r2, size=10, italic=True)
    r3 = p.add_run(source)
    set_font(r3, size=10, color=(100, 116, 139))

# ── Save ──────────────────────────────────────────────────────────────────────
output = 'scoring_system_explanation.docx'
doc.save(output)
print(f'Saved: {output}')
